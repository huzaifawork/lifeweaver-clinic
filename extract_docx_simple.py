#!/usr/bin/env python3
"""
Simple script to extract DOCX file content exactly as it appears
"""

import sys
import os

def install_docx():
    """Install python-docx if not available"""
    try:
        from docx import Document
        return Document
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
        from docx import Document
        return Document

def extract_docx_simple(filename):
    """Extract DOCX file content in the exact order it appears"""
    Document = install_docx()
    
    if not os.path.exists(filename):
        print(f"Error: File '{filename}' not found")
        return
    
    try:
        doc = Document(filename)
        
        print("=== DOCUMENT STRUCTURE ANALYSIS ===")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Total tables: {len(doc.tables)}")
        print()
        
        markdown_content = []
        
        # Process all paragraphs in document order
        print("=== PROCESSING PARAGRAPHS ===")
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:  # Only process non-empty paragraphs
                style = paragraph.style.name if paragraph.style else 'Normal'
                
                print(f"Para {i+1:3d}: [{style:15s}] {text[:80]}...")
                
                # Convert styles to markdown
                if 'Title' in style:
                    markdown_content.append(f"# {text}")
                elif 'Heading 1' in style or style == 'Heading1':
                    markdown_content.append(f"# {text}")
                elif 'Heading 2' in style or style == 'Heading2':
                    markdown_content.append(f"## {text}")
                elif 'Heading 3' in style or style == 'Heading3':
                    markdown_content.append(f"### {text}")
                elif 'Heading 4' in style or style == 'Heading4':
                    markdown_content.append(f"#### {text}")
                elif 'Heading 5' in style or style == 'Heading5':
                    markdown_content.append(f"##### {text}")
                elif 'Heading 6' in style or style == 'Heading6':
                    markdown_content.append(f"###### {text}")
                else:
                    markdown_content.append(text)
                
                markdown_content.append("")  # Add blank line
        
        # Add separator before tables
        if doc.tables:
            markdown_content.append("---")
            markdown_content.append("")
            markdown_content.append("# DOCUMENT TABLES")
            markdown_content.append("")
            
            print("=== PROCESSING TABLES ===")
            for table_idx, table in enumerate(doc.tables):
                print(f"Table {table_idx + 1}: {len(table.rows)} rows, {len(table.columns)} columns")
                
                markdown_content.append(f"## Table {table_idx + 1}")
                markdown_content.append("")
                
                # Extract table data
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell_idx, cell in enumerate(row.cells):
                        # Get all text from cell
                        cell_text = ""
                        for para in cell.paragraphs:
                            if cell_text:
                                cell_text += " "
                            cell_text += para.text.strip()
                        
                        # Clean cell text for markdown
                        cell_text = cell_text.replace("|", "\\|").replace("\n", " ").replace("\r", " ")
                        row_data.append(cell_text.strip())
                    
                    table_data.append(row_data)
                    print(f"  Row {row_idx + 1}: {row_data}")
                
                # Create markdown table
                if table_data:
                    # Ensure all rows have same number of columns
                    max_cols = max(len(row) for row in table_data) if table_data else 0
                    for row in table_data:
                        while len(row) < max_cols:
                            row.append("")
                    
                    if max_cols > 0:
                        # Header row
                        headers = table_data[0]
                        markdown_content.append("| " + " | ".join(headers) + " |")
                        markdown_content.append("| " + " | ".join(["---"] * len(headers)) + " |")
                        
                        # Data rows
                        for row in table_data[1:]:
                            markdown_content.append("| " + " | ".join(row) + " |")
                
                markdown_content.append("")
        
        # Write to file
        output_filename = filename.replace('.docx', '_exact_extraction.md')
        with open(output_filename, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_content))
        
        print(f"\nExact extraction completed: {output_filename}")
        print(f"Total lines in markdown: {len(markdown_content)}")
        
        return output_filename
        
    except Exception as e:
        print(f"Error processing file: {e}")
        import traceback
        traceback.print_exc()
        return None

if __name__ == "__main__":
    result = extract_docx_simple("Internal Doc - Mdm. Clients Name.docx")
    if result:
        print(f"\nSuccess! Check the file: {result}")
