#!/usr/bin/env python3
"""
Script to extract DOCX file content exactly as it appears and convert to Markdown
"""

import sys
import os


def install_docx():
    """Install python-docx if not available"""
    try:
        from docx import Document
        return Document
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call(
            [sys.executable, '-m', 'pip', 'install', 'python-docx'])
        from docx import Document
        return Document


def extract_docx_exact(filename):
    """Extract DOCX file content exactly as it appears"""
    Document = install_docx()

    if not os.path.exists(filename):
        print(f"Error: File '{filename}' not found")
        return

    try:
        doc = Document(filename)

        print("=== DOCUMENT STRUCTURE ANALYSIS ===")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Total tables: {len(doc.tables)}")
        print()

        markdown_content = []

        # Simple approach: process all paragraphs first, then all tables
        # This preserves the content exactly as it appears

        # Process all paragraphs in order
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:  # Only process non-empty paragraphs
                style = paragraph.style.name if paragraph.style else 'Normal'

                print(
                    f"Paragraph {i+1}: Style='{style}', Text='{text[:50]}...'")

                # Convert styles to markdown more accurately
                if 'Title' in style:
                    markdown_content.append(f"# {text}")
                elif 'Heading 1' in style or style == 'Heading1':
                    markdown_content.append(f"# {text}")
                elif 'Heading 2' in style or style == 'Heading2':
                    markdown_content.append(f"## {text}")
                elif 'Heading 3' in style or style == 'Heading3':
                    markdown_content.append(f"### {text}")
                elif 'Heading 4' in style or style == 'Heading4':
                    markdown_content.append(f"#### {text}")
                elif 'Heading 5' in style or style == 'Heading5':
                    markdown_content.append(f"##### {text}")
                elif 'Heading 6' in style or style == 'Heading6':
                    markdown_content.append(f"###### {text}")
                else:
                    markdown_content.append(text)

                # Add blank line after each paragraph
                markdown_content.append("")

        # Process all tables
        if doc.tables:
            markdown_content.append("---")
            markdown_content.append("")
            markdown_content.append("# TABLES SECTION")
            markdown_content.append("")

            for table_idx, table in enumerate(doc.tables):
                print(f"Processing Table {table_idx + 1}")
                markdown_content.append(f"## Table {table_idx + 1}")
                markdown_content.append("")

                # Get table data with exact cell content
                table_data = []
                max_cols = 0

                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell_idx, cell in enumerate(row.cells):
                        # Get all text from cell, including from multiple paragraphs
                        cell_text = ""
                        for para in cell.paragraphs:
                            if cell_text:
                                cell_text += " "  # Use space instead of newline for better table formatting
                            cell_text += para.text.strip()

                        # Clean up cell text for markdown table
                        cell_text = cell_text.replace("|", "\\|").replace(
                            "\n", " ").replace("\r", " ")
                        row_data.append(cell_text.strip())

                    table_data.append(row_data)
                    max_cols = max(max_cols, len(row_data))
                    print(f"  Row {row_idx + 1}: {len(row_data)} columns")

                if table_data and max_cols > 0:
                    # Ensure all rows have the same number of columns
                    for row in table_data:
                        while len(row) < max_cols:
                            row.append("")

                    # Create markdown table
                    if len(table_data) > 0:
                        # First row as header
                        headers = table_data[0]
                        markdown_content.append(
                            "| " + " | ".join(headers) + " |")
                        markdown_content.append(
                            "| " + " | ".join(["---"] * len(headers)) + " |")

                        # Rest of the rows
                        for row in table_data[1:]:
                            markdown_content.append(
                                "| " + " | ".join(row) + " |")

                markdown_content.append("")  # Add space after table

        # Write to markdown file
        output_filename = filename.replace('.docx', '_exact_extraction.md')
        with open(output_filename, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_content))

        print(f"Exact extraction markdown file created: {output_filename}")
        print()
        print("=== CONTENT PREVIEW ===")
        # Show first 50 lines
        for i, line in enumerate(markdown_content[:50]):
            print(f"{i+1:3d}: {line}")

        if len(markdown_content) > 50:
            print(
                f"\n... ({len(markdown_content) - 50} more lines in the file)")

    except Exception as e:
        print(f"Error processing file: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    extract_docx_exact("Internal Doc - Mdm. Clients Name.docx")
